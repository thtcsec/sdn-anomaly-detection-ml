"""
Fix Hình 13 captions + xuất file DIFF dễ Ctrl+F trong Word.
"""
from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "KhoaLuanTotNghiep.docx"
BACKUP = ROOT / "KhoaLuanTotNghiep.backup.docx"
DIFF_OUT = ROOT / "DOCX_CHANGES.txt"


def set_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def fix_captions(doc: Document) -> None:
    new_cap = "Hình 13. Biểu đồ SHAP mean |value| (trung bình các lớp) của mô hình XGBoost"
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t.startswith("Hình 13."):
            continue
        # TOC thường có tab + số trang
        if "\t" in p.text:
            page = p.text.split("\t")[-1].strip()
            # page may itself contain duplicated caption; keep digits only if possible
            digits = "".join(ch for ch in page if ch.isdigit())
            set_para_text(p, f"{new_cap}\t{digits or '55'}")
        else:
            set_para_text(p, new_cap)
        print(f"[✓] Fixed Hình 13 at paragraph {i}")


def write_diff() -> None:
    cur = Document(str(DOCX))
    old = Document(str(BACKUP)) if BACKUP.exists() else None

    lines = []
    lines.append("DOCX CHANGES — đối chiếu KhoaLuanTotNghiep.docx vs backup")
    lines.append("=" * 72)
    lines.append("")
    lines.append("CÁCH XEM TRONG WORD:")
    lines.append("1. Đóng Word hoàn toàn rồi mở lại file KhoaLuanTotNghiep.docx")
    lines.append("2. Ctrl+F tìm đúng các cụm bên dưới (cột SAU)")
    lines.append("3. Mục lục (TOC): click phải Mục lục → Update Field / Cập nhật trường")
    lines.append("   (TOC hay giữ text cũ cho đến khi Update Field)")
    lines.append("")

    checks = [
        ("Phạm vi nghiên cứu (0.4)", "Phạm vi:"),
        ("Thu thập dữ liệu (0.5)", "Thu thập dữ liệu:"),
        ("Tiền xử lý (0.5)", "Tiền xử lý:"),
        ("Dataset synthetic (2.1)", "DDoS synthetic"),
        ("Nhận xét 4.4", "Cả ba mô hình đều đạt hiệu năng trên 98%"),
        ("Hạn chế 4.5", "Phạm vi nhãn còn hẹp"),
        ("CV K=10 (4.7.1)", "K = 10"),
        ("CV 99.96% (4.7.1)", "99.96%"),
        ("Baseline số liệu (4.7.2)", "0.88–0.97"),
        ("SHAP tp_src (4.7.3)", "Theo giá trị mean |SHAP|"),
        ("Hình 13 caption", "Hình 13."),
    ]

    def find_para(doc, needle):
        for p in doc.paragraphs:
            if needle in p.text:
                return p.text.strip()[:220]
        return "(không tìm thấy)"

    for title, needle in checks:
        lines.append(f"## {title}")
        lines.append(f"Tìm (Ctrl+F): {needle}")
        if old is not None:
            lines.append(f"TRƯỚC: {find_para(old, needle) if needle != 'DDoS synthetic' else find_para(old, '506 mẫu DDoS')}")
            # better before needles
        before_map = {
            "Phạm vi:": "Phạm vi:",
            "Thu thập dữ liệu:": "Thu thập dữ liệu:",
            "Tiền xử lý:": "Tiền xử lý:",
            "DDoS synthetic": "506 mẫu DDoS",
            "Cả ba mô hình đều đạt hiệu năng trên 98%": "Isolation Forest là lựa chọn tốt nhất cho unsupervised",
            "Phạm vi nhãn còn hẹp": "Chỉ 3 loại traffic",
            "K = 10": "trực tiếp trên tập huấn luyện",
            "99.96%": "độ lệch chuẩn siêu nhỏ",
            "0.88–0.97": "từ  đến",
            "Theo giá trị mean |SHAP|": "byte_count_per_sec (Tốc độ byte trên giây)",
            "Hình 13.": "Hình 13.",
        }
        if old is not None:
            b_needle = before_map.get(needle, needle)
            lines.append(f"TRƯỚC: {find_para(old, b_needle)}")
        lines.append(f"SAU  : {find_para(cur, needle)}")
        lines.append("")

    # table AE
    lines.append("## Bảng so sánh — hàng Autoencoder")
    if len(cur.tables) > 7:
        ae = [c.text.strip() for c in cur.tables[7].rows[2].cells]
        lines.append(f"SAU  : {ae}")
    if old is not None and len(old.tables) > 7:
        ae_old = [c.text.strip() for c in old.tables[7].rows[2].cells]
        lines.append(f"TRƯỚC: {ae_old}")
    lines.append("")
    lines.append("File backup gốc: KhoaLuanTotNghiep.backup.docx")
    DIFF_OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"[✓] Wrote {DIFF_OUT.name}")


def main() -> None:
    doc = Document(str(DOCX))
    fix_captions(doc)
    doc.save(str(DOCX))
    write_diff()


if __name__ == "__main__":
    main()
