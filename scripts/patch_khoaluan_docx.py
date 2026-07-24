"""
Patch KhoaLuanTotNghiep.docx:
- Sửa phạm vi / phương pháp cho khớp thực tế
- Điền Precision/Recall Autoencoder ở bảng so sánh
- Sửa nhận xét 4.4 (bỏ mâu thuẫn)
- Mở rộng 4.5 Hạn chế
- Điền số liệu trống ở 4.7.2 Baseline
- Cập nhật 4.7.3 SHAP theo kết quả thật
- Thêm ghi chú synthetic DDoS ở 2.1
- Thay hình SHAP (image14) bằng reports/shap_bar.png
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "KhoaLuanTotNghiep.docx"
BACKUP_PATH = ROOT / "KhoaLuanTotNghiep.backup.docx"
SHAP_IMG = ROOT / "reports" / "shap_bar.png"


REPLACEMENTS = {
    # 0.4 Đối tượng / phạm vi
    "Đối tượng: Các dòng lưu lượng mạng (Flow-based) dựa trên giao thức OpenFlow và các thuật toán học máy (Random Forest, XGBoost, Isolation Forest, Autoencoder).":
        "Đối tượng: Các dòng lưu lượng mạng (Flow-based) dựa trên giao thức OpenFlow và các thuật toán học máy (XGBoost, Isolation Forest, Autoencoder).",

    "Phạm vi: Thực nghiệm trên tập dữ liệu InSDN và dữ liệu mô phỏng trong môi trường mạng giả lập bằng Mininet, sử dụng bộ điều khiển Ryu (hoặc ONOS).":
        "Phạm vi: Thực nghiệm trên tập dữ liệu tự thu thập từ môi trường mạng SDN giả lập bằng Mininet, sử dụng bộ điều khiển os-ken (fork của Ryu). Tập dữ liệu công khai InSDN được tham khảo trong phần tổng quan và định hướng phát triển, không dùng làm tập huấn luyện chính.",

    # 0.5 Phương pháp
    "Thu thập dữ liệu: Sử dụng tập dữ liệu chuẩn và tự sinh traffic thực tế.":
        "Thu thập dữ liệu: Tự sinh và thu thập traffic OpenFlow trên testbed Mininet (normal, DDoS, portscan).",

    "Tiền xử lý: Trích xuất các đặc trưng dòng (Flow features) và tính toán các chỉ số thống kê (Entropy).":
        "Tiền xử lý: Trích xuất đặc trưng dòng OpenFlow (10 features), làm sạch, chuẩn hóa StandardScaler và cân bằng SMOTE trên tập train.",

    # 4.4 nhận xét — bỏ câu cũ mâu thuẫn
    "Nhận xét: XGBoost vượt trội khi có nhãn; Isolation Forest là lựa chọn tốt nhất cho unsupervised; Cả ba mô hình đều đạt hiệu năng trên 98%. XGBoost vượt trội tuyệt đối (100%) nhờ khai thác nhãn huấn luyện. Trong nhóm unsupervised, Autoencoder (99.65%) vượt nhẹ Isolation Forest (98.16%) nhờ khả năng học phi tuyến của mạng nơ-ron. Cả hai đều phù hợp cho phát hiện tấn công zero-day không cần nhãn.":
        "Cả ba mô hình đều đạt hiệu năng trên 98%. XGBoost vượt trội tuyệt đối (100%) nhờ khai thác nhãn huấn luyện. Trong nhóm unsupervised, Autoencoder (99.65%) vượt nhẹ Isolation Forest (98.16%) nhờ khả năng học phi tuyến của mạng nơ-ron. Cả hai đều phù hợp cho phát hiện tấn công zero-day không cần nhãn.",

    # 4.5 — thay 2 bullet mỏng bằng đoạn đầy đủ hơn (replace từng dòng riêng)
    "Chỉ 3 loại traffic, thực tế có nhiều biến thể tấn công hơn":
        "Phạm vi nhãn còn hẹp: mới gồm Normal, DDoS và Portscan; chưa bao quát các biến thể tấn công nâng cao (slowloris, low-rate DDoS, exploitation).",

    "Chỉ mới test trên môi trường giả lập.":
        "Thực nghiệm chủ yếu trên Mininet (ít nhiễu nền hơn mạng thực). Một phần mẫu DDoS được bổ sung synthetic để giảm mất cân bằng lớp — cần ghi nhận khi diễn giải độ chính xác cao. Hướng phát triển: thu thêm traffic thật trên lab, đánh giá chéo trên dataset SDN công khai (InSDN) và mở rộng phản ứng tự động.",

    # 4.7.2 — điền số liệu trống
    "Nhóm phương pháp truyền thống (Static Threshold, Multi-Rule IDS, Z-Score): Có điểm Precision khá cao (từ  đến ), tuy nhiên chỉ số Recall và Accuracy lại thấp một cách báo động (Recall chỉ đạt từ  đến ). Điều này cho thấy các luật tĩnh hoặc ngưỡng cố định bị tin tặc qua mặt rất dễ dàng, dẫn đến việc bỏ sót hầu hết các cuộc tấn công động (Tỷ lệ âm tính giả - False Negative cực cao), khiến F1-Score kéo thấp xuống mức dưới .":
        "Nhóm phương pháp truyền thống (Static Threshold, Multi-Rule IDS, Z-Score): Có điểm Precision khá cao (khoảng 0.88–0.97), tuy nhiên Recall và Accuracy thấp rõ rệt (Recall chỉ khoảng 0.024–0.046). Các luật/ngưỡng tĩnh dễ bỏ sót tấn công (False Negative cao), kéo F1-Score xuống dưới 0.09 (baseline tốt nhất khoảng 0.087).",

    "Nhóm mô hình học máy đề xuất (XGBoost, Isolation Forest, Autoencoder): Mang lại hiệu năng áp đảo toàn diện ở cả 4 chỉ số với các cột điểm số đều tiệm cận hoặc đạt mức tuyệt đối  (). Trong đó, XGBoost đạt độ chính xác tuyệt đối . Các mô hình không giám sát như Autoencoder () và Isolation Forest () cũng giữ phong độ cực kỳ ổn định trên .":
        "Nhóm mô hình học máy đề xuất (XGBoost, Isolation Forest, Autoencoder): Hiệu năng vượt trội ở cả 4 chỉ số. XGBoost đạt Accuracy/F1 = 1.0000. Autoencoder đạt Accuracy 0.9965 và F1 0.9982; Isolation Forest đạt Accuracy 0.9816 và F1 0.9904.",

    "Kết luận: Giải pháp học máy XGBoost giúp cải thiện tới +1049% chỉ số F1-Score (tăng từ  lên ) so với phương pháp luật tĩnh tốt nhất. Kết quả thực nghiệm này minh chứng rõ ràng việc chuyển dịch từ quản trị mạng dựa trên luật cấu hình thủ công sang tự động nhận diện bằng học máy là bước đi mang tính quyết định để bảo vệ hạ tầng mạng SDN trước các kịch bản tấn công tinh vi ngày nay.":
        "Kết luận: XGBoost nâng F1-Score từ khoảng 0.087 (baseline tốt nhất) lên 1.0000, tương đương mức cải thiện rất lớn so với ngưỡng tĩnh/luật thủ công. Kết quả khẳng định việc chuyển từ quản trị dựa trên luật cố định sang nhận diện bằng học máy là cần thiết để bảo vệ hạ tầng SDN.",
}


def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Giữ style run đầu, ghi đè toàn bộ text paragraph."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def patch_dataset_section(doc: Document) -> None:
    """Thêm ghi chú synthetic sau mô tả 506 mẫu DDoS."""
    target = "Các mẫu dữ liệu được trích xuất dưới dạng Flow Statistics từ giao thức OpenFlow và sử dụng cho quá trình huấn luyện và đánh giá mô hình học máy."
    new = (
        "Các mẫu dữ liệu được trích xuất dưới dạng Flow Statistics từ giao thức OpenFlow và sử dụng "
        "cho quá trình huấn luyện và đánh giá mô hình học máy. Do lớp DDoS ban đầu thu được trên lab "
        "còn rất ít, nhóm bổ sung thêm mẫu DDoS synthetic (mô phỏng SYN/UDP/ICMP flood dựa trên phân "
        "phối đặc trưng thực nghiệm) để giảm mất cân bằng lớp trước khi áp dụng SMOTE; các kết quả "
        "đánh giá cần được diễn giải trong bối cảnh testbed giả lập này."
    )
    for p in doc.paragraphs:
        if p.text.strip() == target:
            replace_paragraph_text(p, new)
            print("[✓] Updated dataset disclosure (2.1)")
            return
    print("[!] Dataset disclosure paragraph not found")


def patch_shap_section(doc: Document) -> None:
    """Cập nhật nhận xét SHAP theo kết quả chạy thật (tp_src đứng đầu)."""
    mapping = {
        "byte_count_per_sec (Tốc độ byte trên giây): Được xác định là đặc trưng có tầm quan trọng cao nhất tổng thể (Mean |SHAP value| lớn nhất).":
            None,  # handled by full-paragraph match below
    }
    # Full paragraph replacements by startswith anchors
    updates = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("byte_count_per_sec (Tốc độ byte trên giây):"):
            updates.append((
                p,
                "Theo giá trị mean |SHAP| trung bình trên các lớp, thứ tự đóng góp quan trọng nhất là: "
                "tp_src, byte_count_per_sec, packet_count, tp_dst và packet_count_per_sec. "
                "Trong đó tp_src (cổng nguồn) có mean |SHAP| cao nhất, phản ánh hành vi chọn cổng "
                "nguồn khác biệt rõ giữa Portscan/DDoS và lưu lượng bình thường trên testbed.",
            ))
        elif t.startswith("tp_src (Cổng nguồn) và tp_dst (Cổng đích):"):
            updates.append((
                p,
                "tp_src và tp_dst tiếp tục đóng vai trò phân tách mạnh với Portscan (quét nhiều cổng đích) "
                "và một phần DDoS (cổng dịch vụ cố định như 80/443). Đây là tín hiệu giải thích được "
                "vì sao ranh giới 3 lớp trên lab khá rõ, đồng thời cảnh báo rủi ro khi triển khai mạng thật "
                "nơi phân bố cổng phức tạp hơn.",
            ))
        elif t.startswith("packet_count (Tổng số gói tin):"):
            updates.append((
                p,
                "packet_count và byte_count_per_sec bổ sung bằng chứng về cường độ luồng: DDoS thường "
                "đi kèm số gói/tốc độ byte rất cao, trong khi Normal ổn định hơn. Các đặc trưng này "
                "khớp với trực giác vận hành OpenFlow và với baseline ngưỡng tĩnh (nhưng ML tận dụng "
                "được tổ hợp đa đặc trưng thay vì một ngưỡng đơn).",
            ))
        elif t.startswith("Các đặc trưng thời gian và giao thức tĩnh (ip_proto, flow_duration, duration_sec):"):
            updates.append((
                p,
                "Nhóm đặc trưng thời gian/giao thức tĩnh (ip_proto, flow_duration, duration_sec, byte_count) "
                "có |SHAP| thấp hơn rõ rệt, cho thấy mô hình ưu tiên tín hiệu cổng và cường độ luồng hơn "
                "là thời gian sống đơn thuần. Kết quả SHAP giúp giảm tính “hộp đen” và giải thích vì sao "
                "độ chính xác trên lab cao: không gian 10 đặc trưng OpenFlow tạo ranh giới tách lớp khá rõ.",
            ))

    for p, new in updates:
        replace_paragraph_text(p, new)
        print(f"[✓] Updated SHAP para: {new[:60]}...")

    if not updates:
        print("[!] No SHAP commentary paragraphs updated")


def fill_ae_table(doc: Document) -> None:
    """Điền Precision/Recall trống của Autoencoder trong bảng so sánh."""
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 3 or len(table.columns) < 7:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[:4] != ["Mô hình (Model)", "Tiếp cận (Approach)", "Classification", "Accuracy"]:
            continue
        ae = table.rows[2].cells
        if ae[0].text.strip() == "Autoencoder":
            if not ae[4].text.strip():
                ae[4].text = "0.9982"
            if not ae[5].text.strip():
                ae[5].text = "0.9982"
            print(f"[✓] Filled Autoencoder Precision/Recall in table {ti}")
            return
    print("[!] Comparison table with empty AE cells not found")


def replace_shap_image(doc: Document) -> None:
    """Thay media/image14.png (Hình 13 SHAP) bằng shap_bar.png mới."""
    if not SHAP_IMG.exists():
        print(f"[!] Missing {SHAP_IMG}")
        return

    # rId24 -> media/image14.png (đã xác định từ audit)
    target_rel = None
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/image") and rel.target_ref.endswith("image14.png"):
            target_rel = rel
            break
    if target_rel is None:
        print("[!] image14.png relationship not found")
        return

    image_part = target_rel.target_part
    with open(SHAP_IMG, "rb") as f:
        image_part._blob = f.read()
    print(f"[✓] Replaced image14.png with {SHAP_IMG.name}")


def apply_replacements(doc: Document) -> None:
    done = 0
    for p in doc.paragraphs:
        key = p.text.strip()
        if key in REPLACEMENTS:
            replace_paragraph_text(p, REPLACEMENTS[key])
            done += 1
            print(f"[✓] Replaced: {key[:70]}...")
    print(f"[*] Exact replacements applied: {done}/{len(REPLACEMENTS)}")


def main() -> None:
    if not DOCX_PATH.exists():
        raise SystemExit(f"Missing {DOCX_PATH}")

    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"[*] Backup -> {BACKUP_PATH.name}")

    doc = Document(str(DOCX_PATH))
    apply_replacements(doc)
    patch_dataset_section(doc)
    patch_shap_section(doc)
    fill_ae_table(doc)
    replace_shap_image(doc)

    doc.save(str(DOCX_PATH))
    print(f"[✓] Saved {DOCX_PATH.name}")


if __name__ == "__main__":
    main()
