"""Fill remaining blank numbers in CV section + update SHAP caption."""
from docx import Document

DOCX = "KhoaLuanTotNghiep.docx"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    doc = Document(DOCX)

    fixes = {
        470: (
            "Để chứng minh tính ổn định và đảm bảo mô hình XGBoost không rơi vào hiện tượng quá khớp "
            "(Overfitting) do dữ liệu mất cân bằng sau khi xử lý bằng phương pháp SMOTE, nghiên cứu tiến hành "
            "thực nghiệm kiểm định chéo K-Fold phân tầng với K = 10 trực tiếp trên tập dữ liệu đã tiền xử lý. "
            "Kết quả phân tích phân phối hiệu năng qua các vòng lặp được minh họa chi tiết tại hình bên dưới:"
        ),
        474: (
            "Biểu đồ bên trái (10-Fold Cross-Validation Scores): Phân phối điểm số của cả 4 chỉ số cốt lõi bao gồm "
            "Accuracy, F1-Score (macro), Precision (macro), và Recall (macro) đều tập trung sát ngưỡng tuyệt đối "
            "(Accuracy trung bình 99.96%). Khoảng biến thiên của các hộp dữ liệu cực kỳ hẹp với độ lệch chuẩn siêu nhỏ "
            "(±0.06%), chứng tỏ mô hình có độ tin cậy và tính ổn định rất cao, không bị biến động đáng kể khi thay đổi "
            "các tập dữ liệu con trong quá trình xáo trộn kiểm thử."
        ),
        475: (
            "Biểu đồ bên phải (Train vs Test Accuracy per Fold): Đường biểu diễn độ chính xác trên tập huấn luyện "
            "(Train Accuracy - đường màu xanh) và tập kiểm tra (Test Accuracy - đường màu đỏ) gần như trùng khít hoàn toàn "
            "qua cả 10 vòng lặp (Fold 1 đến Fold 10). Độ lệch lớn nhất xuất hiện cục bộ ở Fold 6 nhưng khoảng cách trung bình "
            "(Overfit gap) chỉ khoảng 0.04% (Train ≈ 100%, Test ≈ 99.96%)."
        ),
        476: (
            "Kết luận: Với Accuracy 99.96% ± 0.06% qua 10-fold và overfit gap ≈ 0.04%, kết quả khẳng định mô hình XGBoost "
            "có khả năng tổng quát hóa tốt trên dữ liệu OpenFlow của testbed SDN, không chỉ phụ thuộc vào một lần "
            "train/test split."
        ),
    }

    for idx, text in fixes.items():
        replace_paragraph_text(doc.paragraphs[idx], text)
        print(f"[✓] Fixed paragraph {idx}")

    for p in doc.paragraphs:
        raw = p.text
        if raw.strip().startswith("Hình 13."):
            suffix = raw[raw.find("\t"):] if "\t" in raw else ""
            replace_paragraph_text(
                p,
                "Hình 13. Biểu đồ SHAP mean |value| (trung bình các lớp) của mô hình XGBoost" + suffix,
            )
            print("[✓] Updated Hình 13 caption")
            break

    doc.save(DOCX)
    print("[✓] Saved")


if __name__ == "__main__":
    main()
